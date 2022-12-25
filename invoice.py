import os
import tkinter as tk
from tkinter import ttk
from util import readData, getNextInvoiceNumber, updateCompanyinfo, updateFinancialData
import time
import docx
from datetime import datetime, timedelta
from tkinter import filedialog as fd
from docx.shared import Pt, Inches
import random

class invoiceDialog:
    def __init__(self, root) -> None:
        #information about all companies that are added
        self.companiesData = readData('data/companies.pickle')
        #only names of the companies that are used for drop down list
        self.CompaniesDDOptions = [x['Ime'] for x in self.companiesData]
        #information about all receivers that are added
        self.receiversData = readData('data/receivers.pickle')
        #only receivers names that will be used for dropdown list
        self.receiversDDOptions = [x['Naziv'] for x in self.receiversData]
        #invoice number
        self.invoiceNumber = str(getNextInvoiceNumber(self.companiesData[0]['Invoices']))
        #current year
        self.invoiceYear = str(time.localtime().tm_year)
        #information about all available products
        self.productsData = readData('./data/products.pickle')
        #only names of products for the drop down list
        self.optionProducts = [x['Ime'] for x in self.productsData]
        #currently existing labels of selected products
        self.existingLabels = []

        #these are UI frames
        self.invoiceWindow = tk.Toplevel(root)
        self.invoiceWindow.title('Fakturisanje - Ottoshop')
        self.container = tk.Frame(self.invoiceWindow)
        self.container.pack()
        self.companyFrame = tk.Frame(self.container)
        self.companyFrame.pack(pady=10, fill='both', expand=True)
        self.receiverFrame = tk.Frame(self.container)
        self.receiverFrame.pack(pady=10,  fill='both', expand=True)
        self.invoiceNumberFrame = tk.Frame(self.container)
        self.invoiceNumberFrame.pack(pady=10,  fill='both', expand=True)
        self.productFrame = tk.Frame(self.container)
        self.productFrame.pack(pady=10, padx=15, fill='both', expand=True)
        self.selectedProductsFrame = tk.Frame(self.container)
        self.selectedProductsFrame.pack(pady=10, padx=10, fill='both', expand=True)
        self.invoicePathFrame = tk.Frame(self.container)
        self.invoicePathFrame.pack(pady=10, padx=10, fill='both', expand=True)

        #VARIABLES ONLY ------------------
        #tk string variable about selected company
        self.selectedCompanyInfo = self.companiesData[0]
        self.selectedCompanyName = tk.StringVar()
        self.selectedCompanyName.set(self.companiesData[0]['Ime'])
        #tk string variable about selected receiver
        self.selectedReceiverInfo = self.receiversData[0]
        self.selectedReceiver = tk.StringVar()
        self.selectedReceiver.set(self.receiversData[0]['Naziv'])
        #tk string variable to hold quantity of product thats being added
        self.quantity = tk.StringVar()
        self.quantity.set(0)
        #tk string variable for currently selected product
        self.selectedProducts = []
        self.selectedProduct = tk.StringVar()
        self.selectedProduct.set(self.optionProducts[0])
        #tk string var for invoice path, it will save invoice on this path when print is selected
        self.invoicePath = tk.StringVar()
        self.invoicePath.set(self.companiesData[0]['Path'])
        #tk string variable for invoiceNumber
        self.invoiceNumberVar = tk.StringVar()
        self.invoiceNumberVar.set(self.invoiceNumber + '-' + self.invoiceYear[-2:])

    def saveDocument(self):
        if self.invoicePath.get() != 'none' and len(self.selectedProducts) != 0:
            def setTitleText(paragraph, run):
                paragraph.paragraph_format.space_after = Pt(0)
                run.font.bold = True
                run.font.italic = True
                run.font.size = Pt(20)
                run.font.name = 'Calibri'
            
            def setNormalText(paragraph, run, spaceBefore, spaceAfter, align, size, bold: bool):
                paragraph.paragraph_format.space_before = Pt(spaceBefore)
                paragraph.paragraph_format.space_after = Pt(spaceAfter)
                match align:
                    case 'left':
                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    case 'right':
                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    case 'center':
                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER   
                run.font.size = Pt(size)
                run.font.name = 'Calibri'
                if bold == True:
                    run.font.bold = True

            def getCurrentDate():
                meseci = {
                    1: 'Januar',
                    2: 'Februar',
                    3: 'Mart',
                    4: 'April',
                    5: 'Maj',
                    6: 'Jun',
                    7: 'Jul',
                    8: 'Avgust',
                    9: 'Septembar',
                    10: 'Oktobar',
                    11: 'Novembar',
                    12: 'Decembar'
                }
                date = datetime.now()
                consDate = datetime(date.year, date.month, date.day)
                res = consDate + timedelta(days=30)
                return f"{date.day}-{meseci[date.month]}-{date.year}", f"{res.day}-{meseci[res.month]}-{res.year}"
            
            def setProductInfo(index):
                index = index - 1
                productInfo = {
                    0: index + 1,
                    1: self.selectedProducts[index][0],
                    2: self.selectedProducts[index][1],
                    3: f'{self.selectedProducts[index][2]:,.2f}'.replace(',', ' '),
                    4: f'{self.selectedProducts[index][1] * self.selectedProducts[index][2]:,.2f}'.replace(',', ' ')
                }
                return productInfo
            # Create a new document
            doc = docx.Document()

            # Createing the table for word document header
            table = doc.add_table(rows=6, cols=2)

            # Set the top and bottom padding for all cells in the table
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.space_before = 72
                    cell.paragraphs[0].paragraph_format.space_after = 72
            
            #set the name of the company
            cell = table.cell(0, 0)
            cell.text = self.companiesData[0]['Ime']
            run = cell.paragraphs[0].runs[0]
            setTitleText(cell.paragraphs[0], cell.paragraphs[0].runs[0])

            #set the invoice number
            cell = table.cell(1, 1)
            cell.text = f"Broj fakture: {self.invoiceNumber}-{self.invoiceYear}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'right', 12, False)

            #setting the company id number
            cell = table.cell(1, 0)
            cell.text = f"Maticni broj: {self.companiesData[0]['Maticni']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #set invoice date
            cell = table.cell(2, 1)
            currentDate, dateDue = getCurrentDate() 
            cell.text = f"Datum prometa: {currentDate}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'right', 12, False)
            cell = table.cell(3, 1)
            currentDate, dateDue = getCurrentDate() 
            cell.text = f"Rok placanja: {dateDue}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'right', 12, False)

            #setting company tax number
            cell = table.cell(2, 0)
            cell.text = f"PIB: {self.companiesData[0]['PIB']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #setting the company type id
            cell = table.cell(3, 0)
            cell.text = f"Sifra delatnosti: {self.companiesData[0]['Delatnost']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #setting the company adress
            cell = table.cell(4, 0)
            cell.text = f"Adresa: {self.companiesData[0]['Adresa']}, {self.companiesData[0]['Grad']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #setting the company email
            cell = table.cell(5, 0)
            cell.text = f"Email: {self.companiesData[0]['Email']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #setting the company bank account number
            cell = table.cell(4, 1)
            cell.text = f"Racun: {self.companiesData[0]['Account'].strip()}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'right', 12, False)

            # Createing the table for word document header
            para = doc.add_paragraph("")
            para.paragraph_format.space_before = 10

            table = doc.add_table(rows=5, cols=1)

            #getting receivers data
            receiver = next(x for x in self.receiversData if x['Naziv'] == self.selectedReceiver.get())     

            #add receiver title
            cell = table.cell(0, 0)
            cell.text = "ZA:"
            setTitleText(cell.paragraphs[0], cell.paragraphs[0].runs[0])

            #set receiver name
            cell = table.cell(1, 0)
            cell.text = f"Ime: {receiver['Ime']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #set receiver name
            cell = table.cell(2, 0)
            cell.text = f"Maticni broj: {receiver['Maticni']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #set receiver tax number
            cell = table.cell(3, 0)
            cell.text = f"PIB: {receiver['PIB']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #set receiver Adress
            cell = table.cell(4, 0)
            cell.text = f"Adresa: {receiver['Grad']}, {receiver['Ulica']}"
            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 0, 0, 'left', 12, False)

            #adding an empty line just to make some space between paragraphs
            para = doc.add_paragraph("")
            para.paragraph_format.space_before = 10
            
            table = doc.add_table(rows=len(self.selectedProducts) + 2, cols=5)
            
            table.style = 'Light Grid'
            
            values = {
                0: 'RB',
                1: 'Artikal',
                2: 'Kolicina',
                3: 'Cena',
                4: 'Ukupno'
            }

            width = {
                0: 0.3,
                1: 3.5,
                2: 0.5,
                3: 1.2,
                4: 1.8
            }

            for x in range(len(self.selectedProducts) + 2):
                if x == len(self.selectedProducts) + 1:
                    cell1 = table.rows[x].cells[0]
                    cell2 = table.rows[x].cells[1]
                    cell3 = table.rows[x].cells[2]
                    cell4 = table.rows[x].cells[3]
                    cell5 = table.rows[x].cells[4]
                    cell1.merge(cell2)
                    cell1.merge(cell3)
                    cell1.merge(cell4)
                    cell1.text = 'UKUPNO ZA NAPLATU'
                    totalPrice = 0
                    for product in self.selectedProducts:
                        totalPrice += product[1] * product[2] 
                    setNormalText(cell1.paragraphs[0], cell1.paragraphs[0].runs[0], 5, 5, 'right', 13, True)
                    cell5.text = f'{totalPrice:,.2f} RSD'.replace(',', ' ')
                    setNormalText(cell5.paragraphs[0], cell5.paragraphs[0].runs[0], 5, 5, 'center', 13, True)
                elif x == 0:
                    for y in range(5):
                        cell = table.rows[x].cells[y]
                        cell.paragraphs[0].text = values[y]
                        cell.width = Inches(width[y])
                else:
                    for y in range(5):
                        productInfo = setProductInfo(x)
                        cell = table.rows[x].cells[y]
                        cell.paragraphs[0].text = str(productInfo[y])
                        if y > 1:
                            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 5, 5, 'center', 12, False)
                        else:
                            setNormalText(cell.paragraphs[0], cell.paragraphs[0].runs[0], 5, 5, 'left', 12, False)
                        cell.width = Inches(width[y])
            
            para = doc.add_paragraph("")
            para.paragraph_format.space_before = 10
            
            table = doc.add_table(rows=2, cols=1)
            cell1 = table.cell(0, 0)
            cell3 = table.cell(1, 0)
            cell1.text = 'Fakturu izdao: '
            setNormalText(cell1.paragraphs[0], cell1.paragraphs[0].runs[0], 5, 5, 'left', 12, False)
            cell3.text = '__________________________________'
            setNormalText(cell3.paragraphs[0], cell3.paragraphs[0].runs[0], 5, 5, 'left', 12, False)

            index = next(i for i, sublist in enumerate(self.companiesData) if sublist['Ime'] == self.selectedCompanyName.get())

            companyName = self.companiesData[index]['Ime']
            receiverName = next(x['Ime'] for x in self.receiversData if x['Naziv'] == self.selectedReceiver.get())
            invoiceDate = f"{self.invoiceNumber}_{self.invoiceYear}"

            path = ''

            try:
                if self.invoicePath.get() != self.selectedCompanyInfo['Path']:
                    self.companiesData[index]['Path'] = self.invoicePath.get()

                re = self.invoiceNumberVar.get().split('-')[0]
                self.companiesData[index]['Invoices'] += f',{re}'
                updateCompanyinfo(self.companiesData)

                if os.path.exists(f"data/printedReceipts/{companyName}-{receiverName}-{invoiceDate}.docx"):
                    doc.save(f"data/printedReceipts/{companyName}-{receiverName}-{invoiceDate}({random.randint(0, 1000)}).docx")
                else:
                    doc.save(f"data/printedReceipts/{companyName}-{receiverName}-{invoiceDate}.docx")

                if os.path.exists(f"{self.invoicePath.get()}/{companyName}-{receiverName}-{invoiceDate}.docx"):
                    doc.save(f"{self.invoicePath.get()}/{companyName}-{receiverName}-{invoiceDate}({random.randint(0, 1000)}).docx")
                    path = f"{self.invoicePath.get()}/{companyName}-{receiverName}-{invoiceDate}({random.randint(0, 1000)}).docx"
                else:
                    doc.save(f"{self.invoicePath.get()}/{companyName}-{receiverName}-{invoiceDate}.docx")
                    path = f"{self.invoicePath.get()}/{companyName}-{receiverName}-{invoiceDate}.docx"

            except Exception as err:
                print(err)

            finances = {
                    1: [],
                    2: [],
                    3: [],
                    4: [],
                    5: [],
                    6: [],
                    7: [],
                    8: [],
                    9: [],
                    10: [],
                    11: [],
                    12: [],
                }
            ss = []
            for x in self.selectedProducts:
                ss.append(x[1] * x[2])
            total = sum(ss)
            print(total)

            data = readData('data/financialdata.ottoshop')
            data.setdefault(self.invoiceYear, finances)
            data[self.invoiceYear][datetime.now().month].append(total)
            print(data)
            updateFinancialData(data)

            from docx2pdf import convert
            # Convert the docx document to a PDF file
            convert(path, 'data/printedReceipts/document.pdf')

            # Open the PDF file with the default PDF viewer
            os.startfile(path, 'print')
            self.invoiceWindow.destroy()

    def updateCompanyData(self, company):
        self.selectedCompanyInfo = next(x for x in self.companiesData if x['Ime'] == company)
        self.lastInvoiceNum= str(int(getNextInvoiceNumber(self.selectedCompanyInfo['Invoices'])))
        self.invoiceNumberVar.set(self.invoiceNumber + '-' + self.invoiceYear[-2:])

    def updateReceiverData(self, receiver):
        self.selectedReceiverInfo = next(x for x in self.receiversData if x['Naziv'] == receiver)

    def getFilePath(self):
        file = fd.askdirectory()
        self.invoicePath.set(file)

    def removeProduct(self, product):
        index = [i for i, list in enumerate(self.selectedProducts) if list[0] == product]
        self.selectedProducts[index[0]][4].destroy()
        self.selectedProducts.pop(index[0])

    #when + is clicked to add product to list, this function executes
    def insertProduct(self, frame):
        price = None
        #quantity must be more than 0
        if int(self.quantity.get()) != 0:
            #checking to see if the product being added is not already on the list
            if any(x[0] == self.selectedProduct.get() for x in self.selectedProducts):
                index = [i for i, sublist in enumerate(self.selectedProducts) if sublist[0] == self.selectedProduct.get()]
                self.selectedProducts[index[0]][1] += int(self.quantity.get())
                self.selectedProducts[index[0]][3].configure(text=f"{index[0] + 1}: {self.selectedProducts[index[0]][0]} ({self.selectedProducts[index[0]][2]}) {self.selectedProducts[index[0]][1]}")
            else:
                selectedProductInfo = next(x for x in self.productsData if x['Ime'] == self.selectedProduct.get())
                price = selectedProductInfo['Cena']
                #generating frame for new product
                frm = tk.Frame(self.selectedProductsFrame)
                frm.pack(padx=3, fill='both', expand=True)
                #creating label for product info
                lbl = ttk.Label(frm)
                lbl.pack(side='left')
                #creating button for product removal
                btn = ttk.Button(frm, text="X", width=1, command=lambda: self.removeProduct(selectedProductInfo['Ime']))
                btn.pack(side='left')
                #adding selected product to selected products array and printing label
                self.selectedProducts.append([self.selectedProduct.get(), int(self.quantity.get()), price, lbl, frm])
                lenght = len(self.selectedProducts) - 1
                lbl.configure(text=f"{lenght + 1}: {self.selectedProducts[lenght][0]} ({self.selectedProducts[lenght][2]}) {self.selectedProducts[lenght][1]}")

    def renderDialog(self):
        #generating dropdown list with companies names
        ttk.Label(self.companyFrame, text='Firma:').pack(padx=20, anchor='w')
        companiesDropDown = tk.OptionMenu(self.companyFrame, self.selectedCompanyName, *self.CompaniesDDOptions, command=self.updateCompanyData)
        companiesDropDown.configure(width=30)
        companiesDropDown.pack(padx=20, anchor="w")

        #generating dropdown list with receivers names
        ttk.Label(self.receiverFrame, text='Primalac:').pack(padx=20, anchor='w')
        receiverDropDown = tk.OptionMenu(self.receiverFrame, self.selectedReceiver, *self.receiversDDOptions, command=self.updateReceiverData)
        receiverDropDown.configure(width=30)
        receiverDropDown.pack(padx=20, anchor="w")

        #this is the invoice number field for current invoice
        invoiceNumberFrame = tk.Frame(self.container)
        invoiceNumberFrame.pack(fill='both', expand=True)
        ttk.Label(self.invoiceNumberFrame, text='Broj Fakture:').pack(padx=20, anchor='w')
        invoiceNumberEntry = tk.Entry(self.invoiceNumberFrame, textvariable=self.invoiceNumberVar)
        invoiceNumberEntry.pack(pady=10, padx=22, fill='both', expand=True)

        #products drop down list, as well as options to add it to currently selected products
        productsDropDown = tk.OptionMenu(self.productFrame, self.selectedProduct, *self.optionProducts)
        productsDropDown.configure(width=22)
        productsDropDown.pack(side='left')
        quantityEntry = tk.Entry(self.productFrame, textvariable=self.quantity, width=3)
        quantityEntry.pack(side='left')
        ttk.Button(self.productFrame, text="+", width=3, command=lambda: self.insertProduct(self.selectedProductsFrame)).pack(side='left')

        #button for path choosing and button for printing the document
        ttk.Button(self.invoicePathFrame, width=30, text="Folder", command=self.getFilePath).pack(anchor='center')
        tk.Entry(self.invoicePathFrame, textvariable=self.invoicePath, width=35).pack(anchor='center', pady=10)
        ttk.Button(self.invoicePathFrame, width=30, text="Stampaj", command=self.saveDocument).pack(anchor='center')
