import tkinter as tk
from tkinter import ttk, filedialog
from docx import Document

class BarCodeDialog():
    def __init__(self, root):
        self.BarCodeWindow = tk.Toplevel(root)
        self.BarCodeWindow.geometry('300x915')
        self.BarCodeWindow.title('Stampaj Bar Kodove')
        
        self.container = tk.Frame(self.BarCodeWindow)
        self.container.pack()

        self.fr_SSize = tk.Frame(self.container)
        self.fr_SSize.pack(pady=10, padx=15, fill='both', expand=True)
        self.sv_spath = tk.StringVar(value='None')
        self.sv_ssize = tk.StringVar()

        self.fr_MSize = tk.Frame(self.container)
        self.fr_MSize.pack(pady=10, padx=15, fill='both', expand=True)
        self.sv_mpath = tk.StringVar(value='None')
        self.sv_msize = tk.StringVar()

        self.fr_LSize = tk.Frame(self.container)
        self.fr_LSize.pack(pady=10, padx=15, fill='both', expand=True)
        self.sv_lpath = tk.StringVar(value='None')
        self.sv_lsize = tk.StringVar()

        self.fr_XLSize = tk.Frame(self.container)
        self.fr_XLSize.pack(pady=10, padx=15, fill='both', expand=True)
        self.sv_xlpath = tk.StringVar(value='None')
        self.sv_xlsize = tk.StringVar()

        self.fr_XXLSize = tk.Frame(self.container)
        self.fr_XXLSize.pack(pady=10, padx=15, fill='both', expand=True)
        self.sv_xxlpath = tk.StringVar(value='None')
        self.sv_xxlsize = tk.StringVar()

        self.fr_XXXLSize = tk.Frame(self.container)
        self.fr_XXXLSize.pack(pady=10, padx=15, fill='both', expand=True)
        self.sv_xxxlpath = tk.StringVar(value='None')
        self.sv_xxxlsize = tk.StringVar()

        self.fr_4XLSize = tk.Frame(self.container)
        self.fr_4XLSize.pack(pady=10, padx=15, fill='both', expand=True)
        self.sv_4xlpath = tk.StringVar(value='None')
        self.sv_4xlsize = tk.StringVar()

        self.fr_submitFrame = tk.Frame(self.container)
        self.fr_submitFrame.pack(pady=10, padx=15, fill='both', expand=True)

    def __del__(self):
        del self

    def addPath(self, size):
        match size:
            case 's':
                self.sv_spath.set(filedialog.askopenfilename())
            case 'm':
                self.sv_mpath.set(filedialog.askopenfilename())
            case 'l':
                self.sv_lpath.set(filedialog.askopenfilename())
            case 'xl':
                self.sv_xlpath.set(filedialog.askopenfilename())
            case 'xxl':
                self.sv_xxlpath.set(filedialog.askopenfilename())
            case 'xxxl':
                self.sv_xxxlpath.set(filedialog.askopenfilename())
            case '4xl':
                self.sv_4xlpath.set(filedialog.askopenfilename())

    def Submit(self):
        document = Document('data/sample.docx')
        paragraph = document.paragraphs[0]
        run = paragraph.add_run()
        pic_path = [self.sv_spath.get(), self.sv_mpath.get(), self.sv_lpath.get(), self.sv_xlpath.get(), self.sv_xxlpath.get(), self.sv_xxlpath.get(), self.sv_4xlpath.get()]
        pic_number = [self.sv_ssize.get(), self.sv_msize.get(), self.sv_lsize.get(), self.sv_xlsize.get(), self.sv_xxlsize.get(), self.sv_xxxlsize.get(), self.sv_4xlsize.get()]
        for value in zip(pic_path, pic_number):
            if value[1] != '':
                if int(value[1]) != 0:
                    for x in range(int(value[1])):
                        run.add_picture(value[0])
        document.save('data/barcodetest.docx')
        from docx2pdf import convert
        import os
        convert('data/barcodetest.docx', 'data/barcodes.pdf')
        pdfPath = __file__.replace('barcode.py', 'data\\barcodes.pdf')
        os.system(f'start {pdfPath}')
        self.BarCodeWindow.destroy()
        del self

    def renderDialog(self):
        ttk.Label(self.fr_SSize, text='S Velicina:').pack(padx=0, anchor='w')
        ttk.Button(self.fr_SSize, width=30, text="Putanja", command=lambda: self.addPath('s')).pack(anchor='center')
        ttk.Label(self.fr_SSize, text=self.sv_spath, textvariable=self.sv_spath, width=30).pack(padx=0, anchor='w')
        ttk.Spinbox(self.fr_SSize, from_=0, to=80, textvariable=self.sv_ssize, wrap=True, width=30).pack(pady=10)

        ttk.Label(self.fr_MSize, text='M Velicina:').pack(padx=0, anchor='w')
        ttk.Button(self.fr_MSize, width=30, text="Putanja", command=lambda: self.addPath('m')).pack(anchor='center')
        ttk.Label(self.fr_MSize, text=self.sv_mpath, textvariable=self.sv_mpath, width=30).pack(padx=0, anchor='w')
        ttk.Spinbox(self.fr_MSize, from_=0, to=80, textvariable=self.sv_msize, wrap=True, width=30).pack(pady=10)

        ttk.Label(self.fr_LSize, text='L Velicina:').pack(padx=0, anchor='w')
        ttk.Button(self.fr_LSize, width=30, text="Putanja", command=lambda: self.addPath('l')).pack(anchor='center')
        ttk.Label(self.fr_LSize, text=self.sv_lpath, textvariable=self.sv_lpath, width=30).pack(padx=0, anchor='w')
        ttk.Spinbox(self.fr_LSize, from_=0, to=80, textvariable=self.sv_lsize, wrap=True, width=30).pack(pady=10)

        ttk.Label(self.fr_XLSize, text='XL Velicina:').pack(padx=0, anchor='w')
        ttk.Button(self.fr_XLSize, width=30, text="Putanja", command=lambda: self.addPath('xl')).pack(anchor='center')
        ttk.Label(self.fr_XLSize, text=self.sv_xlpath, textvariable=self.sv_xlpath, width=30).pack(padx=0, anchor='w')
        ttk.Spinbox(self.fr_XLSize, from_=0, to=80, textvariable=self.sv_xlsize, wrap=True, width=30).pack(pady=10)

        ttk.Label(self.fr_XXLSize, text='XXL Velicina:').pack(padx=0, anchor='w')
        ttk.Button(self.fr_XXLSize, width=30, text="Putanja", command=lambda: self.addPath('xxl')).pack(anchor='center')
        ttk.Label(self.fr_XXLSize, text=self.sv_xxlpath, textvariable=self.sv_xxlpath, width=30).pack(padx=0, anchor='w')
        ttk.Spinbox(self.fr_XXLSize, from_=0, to=80, textvariable=self.sv_xxlsize, wrap=True, width=30).pack(pady=10)

        ttk.Label(self.fr_XXXLSize, text='XXXL Velicina:').pack(padx=0, anchor='w')
        ttk.Button(self.fr_XXXLSize, width=30, text="Putanja", command=lambda: self.addPath('xxxl')).pack(anchor='center')
        ttk.Label(self.fr_XXXLSize, text=self.sv_xxxlpath, textvariable=self.sv_xxxlpath, width=30).pack(padx=0, anchor='w')
        ttk.Spinbox(self.fr_XXXLSize, from_=0, to=80, textvariable=self.sv_xxxlsize, wrap=True, width=30).pack(pady=10)

        ttk.Label(self.fr_4XLSize, text='4XL Velicina:').pack(padx=0, anchor='w')
        ttk.Button(self.fr_4XLSize, width=30, text="Putanja", command=lambda: self.addPath('4xl')).pack(anchor='center')
        ttk.Label(self.fr_4XLSize, text=self.sv_4xlpath, textvariable=self.sv_4xlpath, width=30).pack(padx=0, anchor='w')
        ttk.Spinbox(self.fr_4XLSize, from_=0, to=80, textvariable=self.sv_4xlsize, wrap=True, width=30).pack(pady=10)

        ttk.Button(self.fr_submitFrame, width=30, text="Stampaj", command=self.Submit).pack(anchor='center')