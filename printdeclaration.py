import tkinter as tk
from tkinter import ttk, filedialog
from docx import Document

class PrintDeclaration():
    def __init__(self, root):
        self.DeclarationWindow = tk.Toplevel(root)
        self.DeclarationWindow.geometry("300x200")
        self.DeclarationWindow.title("Deklaracije - Ottoshop")

        self.container = tk.Frame(self.DeclarationWindow)
        self.container.pack()

        self.fr_path = tk.Frame(self.container)
        self.fr_path.pack(pady=10, padx=15, fill='both', expand=True)

        self.sv_path = tk.StringVar()
        self.sv_ammount = tk.StringVar()
    
    def __del__(self):
        del self

    def add_path(self):
        self.sv_path.set(filedialog.askopenfilename())
    
    def submit(self):
        document = Document('data/sample.docx')
        paragraph = document.paragraphs[0]
        run = paragraph.add_run()
        for value in range(int(self.sv_ammount.get())):
            run.add_picture(self.sv_path.get())
        document.save('data/barcodetest.docx')
        from docx2pdf import convert
        import os
        convert('data/barcodetest.docx', 'data/barcodes.pdf')
        pdfPath = __file__.replace('barcode.py', 'data\\barcodes.pdf')
        os.system(f'start {pdfPath}')
        self.DeclarationWindow.destroy()
        del self

    def render_dialog(self):
        ttk.Label(self.fr_path, text='S Velicina:').pack(padx=0, anchor='w')
        ttk.Button(self.fr_path, width=30, text="Putanja", command=self.add_path).pack(anchor='center')
        ttk.Label(self.fr_path, text=self.sv_path, textvariable=self.sv_path).pack(padx=0, anchor='w')
        ttk.Spinbox(self.fr_path, from_=0, to=200, textvariable=self.sv_ammount, wrap=True, width=30).pack(pady=10)
        ttk.Button(self.fr_path, width=30, text="Stampaj", command=self.submit).pack(anchor='center')
